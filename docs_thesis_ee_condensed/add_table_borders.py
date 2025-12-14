#!/usr/bin/env python3
"""
Post-process Pandoc-generated DOCX to add table borders, center tables/images, and keep-together properties.
Usage: python add_table_borders.py input.docx output.docx
"""

import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def center_table(table):
    """Center a table horizontally."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)

    tblPr.append(tblBorders)


def keep_table_together(table):
    """Prevent table from splitting across pages."""
    # Set cantSplit on all rows
    for row in table.rows:
        trPr = row._element.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        cantSplit.set(qn("w:val"), "1")
        trPr.append(cantSplit)

    # Set keep-together and keep-with-next on all paragraphs in all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._element.get_or_add_pPr()

                # Keep lines together
                keepLines = OxmlElement("w:keepLines")
                keepLines.set(qn("w:val"), "1")
                pPr.append(keepLines)

                # Keep with next paragraph
                keepNext = OxmlElement("w:keepNext")
                keepNext.set(qn("w:val"), "1")
                pPr.append(keepNext)


def set_keep_together(paragraph):
    """Set paragraph to keep with next paragraph (prevents page breaks)."""
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = OxmlElement("w:keepNext")
    keepNext.set(qn("w:val"), "1")
    pPr.append(keepNext)


def set_page_break_before(paragraph):
    """Add page break before paragraph if needed."""
    pPr = paragraph._element.get_or_add_pPr()
    pageBreakBefore = OxmlElement("w:pageBreakBefore")
    pageBreakBefore.set(qn("w:val"), "1")
    pPr.append(pageBreakBefore)


def main():
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} input.docx output.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    print(f"Reading {input_file}...")
    doc = Document(input_file)

    table_count = 0
    image_count = 0

    # Process each element in the document body
    for i, element in enumerate(doc.element.body):
        # Check if this is a table
        if element.tag.endswith("}tbl"):
            # Find the table object
            for table in doc.tables:
                if table._element == element:
                    # Add borders
                    add_table_borders(table)

                    # Center the table
                    center_table(table)

                    # Keep all rows together (prevent table from splitting)
                    keep_table_together(table)

                    table_count += 1

                    # Find the paragraph immediately before the table (likely a caption)
                    # Look at the 2 paragraphs before in case there's a spacer
                    for j in range(max(0, i - 2), i):
                        check_element = doc.element.body[j]
                        if check_element.tag.endswith("}p"):
                            # Keep this paragraph with the next (table)
                            for para in doc.paragraphs:
                                if para._element == check_element:
                                    set_keep_together(para)
                                    break

                    break

    # Center all paragraphs containing images (charts/diagrams)
    for paragraph in doc.paragraphs:
        # Check if paragraph contains an image
        if any(run._element.xpath(".//pic:pic") for run in paragraph.runs):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_count += 1

    print(f"Added borders and keep-together properties to {table_count} tables")
    print(f"Centered {table_count} tables and {image_count} images")
    print(f"Writing {output_file}...")
    doc.save(output_file)
    print("Done!")


if __name__ == "__main__":
    main()
